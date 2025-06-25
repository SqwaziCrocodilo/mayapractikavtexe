# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
import sqlite3

def set_font(cell, bold=False, size=8, align_left=False):
    """Настройка шрифта и выравнивания"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.name = 'Arial'
        if align_left:
            paragraph.alignment = 0  # По левому краю
        else:
            paragraph.alignment = 1  # Центрирование

def create_table(doc):
    """Создание таблицы с 8 столбцами"""
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    
    # Ширина столбцов (в дюймах)
    widths = [
        Inches(0.7),  # Код блюда
        Inches(3.5),  # Наименование
        Inches(0.7),  # Кол-во порций
        Inches(0.8),  # НЕТТО (порция)
        Inches(0.8),  # БРУТТО (порция)
        Inches(0.8),  # НЕТТО (общее)
        Inches(0.8),  # БРУТТО (общее)
        Inches(0.8),  # Выход (гр)
    ]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Основные заголовки
    hdr = table.rows[0].cells
    hdr[0].text = "Код\nблюда"
    hdr[1].text = "Наименование"
    hdr[2].text = "Кол-во\nпорций"
    hdr[3].merge(hdr[4])
    hdr[3].text = "Кол-во на порцию (гр)"
    hdr[5].merge(hdr[6])
    hdr[5].text = "Общее кол-во (кг)"
    hdr[7].text = "Выход\n(гр)"

    # Подзаголовки
    sub_hdr = table.rows[1].cells
    sub_hdr[3].text = "НЕТТО"
    sub_hdr[4].text = "БРУТТО"
    sub_hdr[5].text = "НЕТТО"
    sub_hdr[6].text = "БРУТТО"
    sub_hdr[7].text = ""

    # Форматирование заголовков
    for cell in table.rows[0].cells + table.rows[1].cells:
        set_font(cell, bold=True)
        cell.paragraphs[0].alignment = 1

    return table

def fetch_dishes(db_path):
    """Получение списка блюд из БД"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT Код_блюда, Наименование, Общий FROM Блюда")
    dishes = cursor.fetchall()
    conn.close()
    return dishes

def fetch_ingredients(db_path, dish_code):
    """Получение ингредиентов"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT Ингридиент, НЕТТО, БРУТТО FROM Ингридиенты WHERE Код_блюда = ?",
        (dish_code,)
    )
    ingredients = cursor.fetchall()
    conn.close()
    return ingredients

def add_dish(doc, db_path, dish_code, portions):
    """Добавление блюда в таблицу"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute(
        "SELECT Наименование, Общий FROM Блюда WHERE Код_блюда = ?",
        (dish_code,)
    )
    dish_name, output = cursor.fetchone()
    ingredients = fetch_ingredients(db_path, dish_code)
    
    table = doc.tables[0]
    row = table.add_row().cells

    # Основные данные
    row[0].text = dish_code
    row[2].text = str(portions)
    row[7].text = str(output)

    # Название блюда и ингредиенты
    dish_text = f"**{dish_name}**"
    for name, _, _ in ingredients:
        dish_text += f"\n{name}"
    row[1].text = dish_text.strip()
    set_font(row[1], align_left=True)

    # Веса ингредиентов
    netto_vals = []
    brutto_vals = []
    total_netto = []
    total_brutto = []
    
    for name, netto, brutto in ingredients:
        netto_clean = netto.replace(',', '.')
        brutto_clean = brutto.replace(',', '.')
        
        netto_val = float(netto_clean)
        brutto_val = float(brutto_clean)
        
        formatted_netto = f"{netto_val:.3f}".replace('.', ',')
        formatted_brutto = f"{brutto_val:.3f}".replace('.', ',')
        total_netto_val = f"{(netto_val * portions / 1000):.3f}".replace('.', ',')
        total_brutto_val = f"{(brutto_val * portions / 1000):.3f}".replace('.', ',')
        
        netto_vals.append(formatted_netto)
        brutto_vals.append(formatted_brutto)
        total_netto.append(total_netto_val)
        total_brutto.append(total_brutto_val)

    # Заполнение данных
    row[3].text = "\n".join(netto_vals)
    row[4].text = "\n".join(brutto_vals)
    row[5].text = "\n".join(total_netto)
    row[6].text = "\n".join(total_brutto)

    # Выравнивание
    for idx in [0, 2, 3, 4, 5, 6, 7]:
        set_font(row[idx])
        row[idx].paragraphs[0].alignment = 1  # Центрирование

    conn.close()

def main():
    db_path = r"D:\Прога по работе\База данных\MENU.db"
    doc = Document()
    
    # Настройка полей
    section = doc.sections[0]
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    
    create_table(doc)
    
    dishes = fetch_dishes(db_path)
    print("Доступные блюда:")
    for idx, (code, name, _) in enumerate(dishes, 1):
        print(f"{idx}. {code} - {name}")
    
    # Ввод нескольких блюд
    while True:
        choices = input("\nВведите номера блюд через запятую (0 - выход): ").strip()
        if choices == "0":
            break
        try:
            selected_indices = [int(x.strip()) - 1 for x in choices.split(',')]
            valid_indices = [idx for idx in selected_indices if 0 <= idx < len(dishes)]
            
            if not valid_indices:
                print("Ошибка: введены недопустимые номера.")
                continue
            
            for idx in valid_indices:
                dish_code = dishes[idx][0]
                portions = int(input(f"Количество порций для {dish_code}: "))
                add_dish(doc, db_path, dish_code, portions)
            
            another = input("Добавить ещё блюда? (да/нет): ").strip().lower()
            if another != "да":
                break
        except ValueError:
            print("Ошибка: введите числа через запятую (например: 1,2,3).")
    
    doc.save("Меню_готовое.docx")
    print("Документ сохранен!")

if __name__ == "__main__":
    main()